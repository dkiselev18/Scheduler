from docx import Document
from ..models import Events


class Report:

    def __init__(self, file_name, day, month, year):
        self.file_name = file_name
        self.document = Document('D:\\Python\\Scheduler\\core\\doc\\day_plan.docx')
        self.year = year
        self.month = month
        self.day = day

    def report_maker(self):
        self._change_table()
        return self.document

    def _change_table(self):
        events = Events.objects.filter(StartDate__day=self.day, StartDate__month=self.month, StartDate__year=self.year)
        table = self.document.tables[0]
        i = 1
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(str(i))
        row.cells[1].paragraphs[0].add_run("Подъем Государственного флага.")
        row.cells[2].paragraphs[0].add_run("8.40-8.55")
        i += 1
        for event in events:
            row = table.add_row()
            row.cells[0].paragraphs[0].add_run(str(i))
            row.cells[1].paragraphs[0].add_run(event.Name)
            i += 1
        return

